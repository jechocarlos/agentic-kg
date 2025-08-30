"""
AKG Document Parser - Hybrid approach using native parsers for text files
and LlamaParse for multimodal documents.
"""
import asyncio
import logging
import os
import ssl
from pathlib import Path
from typing import Any, Dict, Optional

from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from ..config import config

logger = logging.getLogger(__name__)


class SSLContextManager:
    """Manages SSL context for LlamaParse API calls."""
    
    @staticmethod
    def create_ssl_context():
        """Create SSL context with proper certificate handling."""
        try:
            # Create SSL context with default settings
            context = ssl.create_default_context()
            
            # For development/testing, you might need to disable certificate verification
            # context.check_hostname = False
            # context.verify_mode = ssl.CERT_NONE
            
            return context
        except Exception as e:
            logger.warning(f"Failed to create SSL context: {e}")
            return None


class DocumentParser:
    """Hybrid document parser using native parsers for text files and LlamaParse for multimodal documents."""
    
    def __init__(self):
        """Initialize the document parser."""
        self.supported_text_formats = {'.md', '.txt', '.csv', '.json', '.xml', '.html', '.py', '.js', '.css', '.yaml', '.yml'}
        self.native_pdf_support = True
        self.multimodal_formats = {'.pptx', '.xlsx', '.docx'}  # Files that may contain images/complex layouts
        
        # Try to initialize LlamaParse for multimodal documents
        self.llamaparse_available = self._init_llamaparse()

    def _init_llamaparse(self) -> bool:
        """Initialize LlamaParse if available and working."""
        try:
            from llama_parse import LlamaParse

            # Test if we can create a LlamaParse instance
            if not config.LLAMAPARSE_API_KEY:
                logger.warning("LlamaParse API key not found - multimodal parsing will be limited")
                return False
                
            # Create SSL context
            ssl_context = SSLContextManager.create_ssl_context()
            if not ssl_context:
                logger.warning("SSL context creation failed - LlamaParse may not work")
                
            self.llamaparse = LlamaParse(api_key=config.LLAMAPARSE_API_KEY)
            logger.info("✅ LlamaParse initialized successfully")
            return True
            
        except ImportError:
            logger.info("LlamaParse not available - using native parsers only")
            return False
        except Exception as e:
            logger.warning(f"LlamaParse initialization failed: {e}")
            return False

    async def parse_document(self, file_path: Path) -> str:
        """Parse document content based on file type and complexity."""
        logger.info(f"🔍 Analyzing document: {file_path.name}")
        
        file_extension = file_path.suffix.lower()
        
        # Route to appropriate parser based on file type
        if file_extension in self.supported_text_formats:
            return await self._parse_text_native(file_path)
        elif file_extension == '.pdf':
            return await self._parse_pdf_native(file_path)
        elif file_extension in self.multimodal_formats:
            return await self._parse_multimodal(file_path)
        else:
            logger.warning(f"⚠️ Unsupported file type: {file_extension}")
            return await self._parse_fallback(file_path)

    async def _parse_text_native(self, file_path: Path) -> str:
        """Parse text-based files using native Python."""
        logger.info(f"📝 Parsing text file {file_path.name} with native Python")
        
        try:
            # Read file with encoding detection
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not decode file with any supported encoding")
            
            # Add file metadata
            metadata_header = f"""# {file_path.name}

**Source:** {file_path}
**Type:** {file_path.suffix.upper()[1:]} Document
**Size:** {len(content)} characters

---

"""
            
            full_content = metadata_header + content
            logger.info(f"✅ Extracted {len(content)} characters from text file")
            return full_content
            
        except Exception as e:
            logger.error(f"❌ Failed to parse text file {file_path}: {e}")
            return await self._parse_fallback(file_path)

    async def _parse_pdf_native(self, file_path: Path) -> str:
        """Parse PDF using native Python libraries (pdfplumber and PyPDF2)."""
        logger.info(f"📖 Parsing PDF {file_path.name} with native Python parsers")
        
        # Try pdfplumber first (better for text extraction and tables)
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                content = []
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        content.append(f"--- Page {page_num} ---\n{text}")
                    
                    # Extract tables if any
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table:
                            table_text = "\n".join([" | ".join(str(cell) if cell else "") for row in table if row for cell in row])
                            content.append(f"\n--- Table {table_num} on Page {page_num} ---\n{table_text}")
                
                if content:
                    full_content = "\n\n".join(content)
                    logger.info(f"✅ Extracted {len(full_content)} characters from PDF using pdfplumber")
                    return full_content
                    
        except ImportError:
            logger.info("pdfplumber not available, trying PyPDF2...")
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber failed for {file_path}: {e}")
            
        # Fallback to PyPDF2
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = []
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        content.append(f"--- Page {page_num} ---\n{text}")
                
                if content:
                    full_content = "\n\n".join(content)
                    logger.info(f"✅ Extracted {len(full_content)} characters from PDF using PyPDF2")
                    return full_content
                    
        except ImportError:
            logger.info("PyPDF2 not available")
        except Exception as e:
            logger.warning(f"⚠️ PyPDF2 failed for {file_path}: {e}")
        
        # Final fallback - return basic info
        logger.info(f"⚠️ PDF appears to be image-based or encrypted - no text content extracted")
        file_stats = file_path.stat()
        return f"""# PDF Document: {file_path.name}

**Source:** {file_path}
**Type:** PDF Document
**Size:** {file_stats.st_size} bytes
**Pages:** Available but content not extractable
**Status:** Image-based or encrypted PDF

This PDF document appears to be image-based (scanned) or encrypted.
Text extraction with native Python libraries was unsuccessful.
Consider using OCR tools for image-based PDFs or providing the password for encrypted PDFs."""

    async def _parse_multimodal(self, file_path: Path) -> str:
        """Parse multimodal documents (PPTX, XLSX, DOCX) using LlamaParse or native parsers."""
        logger.info(f"🎨 Parsing multimodal document {file_path.name}")
        
        # Try LlamaParse first for best results with multimodal content
        if self.llamaparse_available:
            try:
                logger.info("Attempting LlamaParse for multimodal document...")
                
                # Use asyncio to handle the parsing
                documents = await asyncio.to_thread(
                    self.llamaparse.load_data, 
                    str(file_path)
                )
                
                if documents:
                    content = "\n\n".join([doc.text for doc in documents])
                    logger.info(f"✅ LlamaParse extracted {len(content)} characters")
                    return content
                    
            except Exception as e:
                logger.warning(f"⚠️ LlamaParse failed for {file_path}: {e}")
        
        # Fallback to native parsing for specific formats
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.docx':
            return await self._parse_docx_native(file_path)
        else:
            logger.warning(f"⚠️ No native parser available for {file_extension}")
            return await self._parse_fallback(file_path)

    async def _parse_docx_native(self, file_path: Path) -> str:
        """Parse DOCX files using python-docx."""
        try:
            doc = DocxDocument(str(file_path))  # Convert Path to string
            
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content.append(row_text)
            
            full_content = "\n\n".join(content)
            logger.info(f"✅ Extracted {len(full_content)} characters from DOCX")
            return full_content
            
        except Exception as e:
            logger.error(f"❌ Failed to parse DOCX {file_path}: {e}")
            return await self._parse_fallback(file_path)

    def get_document_metadata(self, file_path: Path, content: Optional[str] = None) -> Dict[str, Any]:
        """Extract metadata from a document file and its content."""
        try:
            file_stats = file_path.stat()
            
            metadata = {
                'file_name': file_path.name,
                'file_size': file_stats.st_size,
                'file_extension': file_path.suffix.lower(),
                'source_path': str(file_path),
                'modified_time': file_stats.st_mtime,
                'file_type': self._get_file_type(file_path),
                'parser_used': self._determine_parser_type(file_path)
            }
            
            # Add content-based metadata if content is provided
            if content:
                metadata.update({
                    'content_length': len(content),
                    'content_preview': content[:200] + '...' if len(content) > 200 else content,
                    'word_count': len(content.split()) if content else 0,
                    'line_count': content.count('\n') + 1 if content else 0
                })
            
            return metadata
            
        except Exception as e:
            logger.error(f"❌ Failed to extract metadata from {file_path}: {e}")
            return {
                'file_name': file_path.name,
                'source_path': str(file_path),
                'error': str(e)
            }
    
    def _get_file_type(self, file_path: Path) -> str:
        """Determine the file type category."""
        extension = file_path.suffix.lower()
        
        if extension in self.supported_text_formats:
            return 'text'
        elif extension == '.pdf':
            return 'pdf'
        elif extension in self.multimodal_formats:
            return 'multimodal'
        else:
            return 'unknown'
    
    def _determine_parser_type(self, file_path: Path) -> str:
        """Determine which parser would be used for this file."""
        extension = file_path.suffix.lower()
        
        if extension in self.supported_text_formats:
            return 'native_text'
        elif extension == '.pdf':
            return 'native_pdf'
        elif extension in self.multimodal_formats:
            return 'llamaparse' if self.llamaparse_available else 'native_limited'
        else:
            return 'fallback'

    async def _parse_fallback(self, file_path: Path) -> str:
        """Fallback parser that returns basic file information."""
        logger.info(f"📋 Using fallback parser for {file_path.name}")
        
        try:
            file_stats = file_path.stat()
            
            return f"""# {file_path.name}

**Source:** {file_path}
**Type:** {file_path.suffix.upper()[1:] if file_path.suffix else 'Unknown'} Document
**Size:** {file_stats.st_size} bytes
**Modified:** {file_stats.st_mtime}
**Status:** Content extraction not available

This document exists but could not be parsed with available tools.
Manual review may be required to extract meaningful content.
"""
        except Exception as e:
            logger.error(f"❌ Even fallback parsing failed for {file_path}: {e}")
            return f"""# {file_path.name}

**Source:** {file_path}
**Status:** File could not be accessed

Error occurred while trying to read this file: {str(e)}
"""
